#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Thu Jul  9 09:39:23 2020

Autmated web scrapper to generate sections for a monthly marketing email.
Takes in codes for Viewpoints, SoC, Patterns and collects title, descp, url.
Combines items together in a document with correct formating including links.
@author: BreakfastMan
"""

# Select system 
system = 'mac'
#system = 'windows'

# Select path for driver
driver_location = '/Users/Breakfastman/Dropbox/Python/Web_Scraping/SBI_Web_Scrape/'

# What monthly group is this for?
group = '2020-02'

# Input SoC, Pattern, and VPs codes for each list
# IT and Digitalization
IT_list = ['SoC1137', '1bgd', '2ct', 'P1460']
# Life Sciences and Health-Care
health_list = ['2bs','1bc','1nb','1ui']
# Energy and Environment
energy_list = ['SoC1142','1es','1fc','1ret']
# Materials and Manufacturing
materials_list = ['23dp','1bp','1ep','1sm']
# Sensors and Electronics
sensors_list = ['1oe','2sns','2iot','2ne']
# Connected Lifestyles
connected_list = ['2ai','1iot','SoC1138','1mc']

# %% Setup web driver and sign in to SBI
from selenium import webdriver
from bs4 import BeautifulSoup

# Import credential file
import SBI_credentials

def start_web(driver_location, system):
    if system == 'mac':
        driver_sys = 'chromedriver_mac'
    elif system == 'windows':
        driver_sys = 'chromedriver_windows'
    else:
        print('System not defined. Must set system variable to mac or windows')
    driver = webdriver.Chrome(driver_location + driver_sys)
    driver.get ('http://www.strategicbusinessinsights.com/login.asp')
    driver.find_element_by_id('USERNAME').send_keys(SBI_credentials.username)
    driver.find_element_by_id('PASSWORD').send_keys(SBI_credentials.passphrase)
    driver.find_element_by_id('clntlogin').click()
    return (driver)

# Starts web driver navigates to SBI website and signs in    
driver = start_web(driver_location, system)

# %% Grabs info from web

def info_fetch(driver, item_list, group):
    base_url = 'http://www.strategicbusinessinsights.com'
    info_dump = []
    
    for item in item_list:    
        if item[0] is 'P':
            # Get info for a Pattern
            # Generate url and navigate to webpage
            url = base_url + '/scan/patterns/' + item + '.shtml'
            driver.get(url)
            
            # Grab content from the webpage
            content = driver.page_source
            soup = BeautifulSoup(content,features='lxml')
            
            # Look through html to find title, text, link name
            title = soup.find(attrs={'class':'intro'}).h1.get_text().split('\n')[0]
            text = soup.find(attrs={'class':'pub-copy-teaser'}).get_text().split('\n')[1]
            link_name = 'Scan Pattern ' + item
            
        elif item[0] is 'S':
            # Get info from a SoC
            # Generate url and navigate to webpage
            url = base_url + '/scan/SoC/' + item + '.shtml'
            url_synop = base_url + '/scan/mtgsynopses/' + group + '.shtml'
            driver.get(url_synop)
            
            # Grab content from webpage
            content = driver.page_source
            soup = BeautifulSoup(content,features='lxml')
            
            # Look through html to find title, text, link name
            title = soup.find(id=item).get_text().split('—')[1][1:]
            text = soup.find(id=item).next_sibling.next_sibling.get_text()
            link_name = 'Scan ' + item
            
        elif item[0] is '1' or item[0] is '2':
            # Get info from VPs
            # Generate URL and navigate to webpage
            tech = item[1:]
            article = int(item[0])
            url = base_url + '/explorer/' + tech + '/' + tech + "-" + group + '.shtml'
            if article == 2:
                url = url + '#2'        
            driver.get(url)
            
            # Grabe content from webpage
            content = driver.page_source
            soup = BeautifulSoup(content,features='lxml')
            
            # Look through html to find title, text, link name
            title = soup.findAll(attrs={'class':'vpts-va-title'})[article-1].get_text()
            text = soup.findAll(attrs={'class':'significance'})[article-1].get_text().split('\n')[2]
            tech_name = soup.find(attrs={'class':'intro'}).h1.get_text().split('\n')[0]
            # Exception for Renewable Energy Technologies
            if tech == 'ret':
                link_name = 'Renewable Energy Technologies'
            else:
                link_name = 'Explorer Technology ' + tech_name
            
        else:
            print('Error, check formating of item list')
            
        def em_dash_replace(text):
            # funciton for replacing em dashes with double dashes
            return(text.replace('—', '--'))
            
        info_dump.append([em_dash_replace(title),em_dash_replace(text),url,link_name])
    return(info_dump)

# Running fetch info for each list
IT_info = info_fetch(driver, IT_list, group)
health_info = info_fetch(driver, health_list, group)
energy_info = info_fetch(driver, energy_list, group)
materials_info = info_fetch(driver, materials_list, group)
sensors_info = info_fetch(driver, sensors_list, group)
connected_info = info_fetch(driver, connected_list, group)
# %% Create functions for outputing items to the document

import docx

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # Applying the grey color to the URL link.
    r.font.color.rgb = docx.shared.RGBColor(0x5C, 0x5C, 0x5C) 
    r.font.size = docx.shared.Pt(10.5)
    r.font.name = 'Arial'
    r.font.underline = True

    return hyperlink

def write_p(document,item):
    p = document.add_paragraph()
    p.style.font.name = 'Arial'
    p.style.font.size = docx.shared.Pt(10.5)
    p.style.font.color.rgb = docx.shared.RGBColor(0x5C, 0x5C, 0x5C)
    
    run_title = p.add_run()
    title_style = run_title.font
    title_style.underline = True
    title_style.bold = True


    run_title.add_text(item[0])
    p.add_run().add_break()
    p.add_run().add_text(item[1] + ' ')
    
    add_hyperlink(p,item[3],item[2])
    
    p.add_run().add_text('.')

# %% Initalize output document
document = docx.Document()
# %% Creating document by generating a paragraph for each list
h = document.add_paragraph()
h.add_run('IT and Digitalization')

for item in IT_info:
    write_p(document,item)
    
h = document.add_paragraph()
h.add_run().add_break()
h.add_run('Life Sciences and Health-Care')


for item in health_info:
    write_p(document,item)

h = document.add_paragraph()
h.add_run().add_break()
h.add_run('Energy and Environment')


for item in energy_info:
    write_p(document,item)
    
h = document.add_paragraph()
h.add_run().add_break()
h.add_run('Materials and Manufacturing')


for item in materials_info:
    write_p(document,item)

h = document.add_paragraph()
h.add_run().add_break()
h.add_run('Sensors and Electronics')


for item in sensors_info:
    write_p(document,item)

h = document.add_paragraph()
h.add_run().add_break()
h.add_run('Connected Lifestyles')


for item in connected_info:
    write_p(document,item)
# %% Saving document
document.save('monthly_email.docx')