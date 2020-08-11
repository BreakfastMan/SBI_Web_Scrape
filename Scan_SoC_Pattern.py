#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Thu Jul  9 09:39:23 2020

Autmated web scrapper to generate sections for Scan reports.
Takes in codes for SoCs and Patterns and collects title, descp, url.
Combines items together in a document with correct formating including links.
@author: BreakfastMan
"""

# Select system 
system = 'mac'
#system = 'windows'

# Select path for driver
driver_location = '/Users/Breakfastman/Dropbox/Python/Web_Scraping/SBI_Web_Scrape/'


# List of SoCs
SoCs = [1156, 1155, 1153, 1151, 1147, 1137, 1116, 1115, 1103]
# List of Patterns
Patterns = [1495, 1494, 1493, 1489, 1482, 1481, 1479, 1470, 1454, 1453, 1446, 
            1445, 1440, 1438, 1429, 1428, 1419, 1405, 1384, 1374, 1368]


# %% Setup web driver and sign in to SBI
from selenium import webdriver
from bs4 import BeautifulSoup
import re
# import credential file
import SBI_credentials

def start_web(driver_location, system):
    if system == 'mac':
        driver_sys = 'chromedriver_mac'
    elif system == 'windows':
        driver_sys = 'chromedriver_windows'
    else:
        print('System not defined. Must set system variable to mac or windows')
    driver = webdriver.Chrome(driver_location + driver_sys)
    driver.get ('http://www.strategicbusinessinsights.com/login.asp')
    driver.find_element_by_id('USERNAME').send_keys(SBI_credentials.username)
    driver.find_element_by_id('PASSWORD').send_keys(SBI_credentials.passphrase)
    driver.find_element_by_id('clntlogin').click()
    return (driver)

# Starts web driver navigates to SBI website and signs in      
driver = start_web(driver_location, system)

# %% Navigate to main Scan page
Scan_url = 'http://www.strategicbusinessinsights.com/scan/signals.shtml'
driver.get(Scan_url)
driver.find_element_by_link_text('Expand all').click()
content = driver.page_source
soup = BeautifulSoup(content, features='lxml')

# %% Prep lists to store data and get synopis group
SoC_info = []
Pattern_info = []

# Data for each item [code, synop group, title, text, url, link_name]
for soc in SoCs:
    soc_item = 'SoC' + str(soc)
    synop_group = soup.find(text=re.compile(soc_item)).find_parent(
        'ul').find_previous_sibling('h5',
                attrs={'class':'pub-item-subitems-secure'}).a['href'][-11:-4]
    SoC_info.append([soc_item, synop_group, None, None, None, None])
    
for pat in Patterns:
    pat_item = 'P' + str(pat)
    synop_group = soup.find(text=re.compile(pat_item)).find_parent(
        'ul').find_previous_sibling('h5',
                attrs={'class':'pub-item-subitems-secure'}).a['href'][-11:-4]
    Pattern_info.append([pat_item, synop_group, None, None, None, None])

# %% Function for scrapping info from webpages
    
def info_fetch(driver, item_list):
    base_url = 'http://www.strategicbusinessinsights.com'

    for item in item_list:    
        if item[0][0] is 'P':
            # Get info for a Pattern
            # Generate url and navigate to webpage
            url = base_url + '/scan/patterns/' + item[0] + '.shtml'
            driver.get(url)
            
            # Grab content from the webpage
            content = driver.page_source
            soup = BeautifulSoup(content,features='lxml')
            
            # Look through html to find title, text, link name
            title = soup.find(attrs={'class':'intro'}).h1.get_text().split('\n')[0]
            text = soup.find(attrs={'class':'pub-copy-teaser'}).get_text().split('\n')[1]
            link_name = 'Scan Pattern ' + item[0]
            
        elif item[0][0] is 'S':
            # Get info for a SoC
            # Generate url and navigate to webpage
            url = base_url + '/scan/SoC/' + item[0] + '.shtml'
            url_synop = base_url + '/scan/mtgsynopses/' + item[1] + '.shtml'
            driver.get(url_synop)
            
            # Grab content from the webpage
            content = driver.page_source
            soup = BeautifulSoup(content,features='lxml')
            
            # Look through html to find title, text, link name
            title = soup.find(id=item[0]).get_text().split('—')[1][1:]
            text = soup.find(id=item[0]).next_sibling.next_sibling.get_text()
            link_name = 'Scan ' + item[0]
            
        else:
            print('Error, check formating of item list')
            
        item[2] = title
        item[3] = text
        item[4] = url
        item[5] = link_name
    return(item_list)

# Run function for SoC and Pattern list
info_fetch(driver, SoC_info)
info_fetch(driver, Pattern_info)
# %% Create functions for outputing items to the document

import docx

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    return hyperlink

def write_p(document,item):
    p = document.add_paragraph()
    p_format = p.paragraph_format
    p_format.first_line_indent = docx.shared.Inches(-.19)
    p_format.left_indent = docx.shared.Inches(.19)
    p_format.space_after = docx.shared.Inches(0)
    p_format.line_spacing = docx.shared.Pt(14)
    p.style.font.name = 'Times'
    p.style.font.size = docx.shared.Pt(12)
    p.style = 'List Bullet'
    
    add_hyperlink(p,item[0] + ' — ' + item[2],item[4])
    
    run_body = p.add_run()
    run_body.add_text(': ' + item[3])

# %% Initalize output document
document = docx.Document()
# %% Write each item in SoC and Pattern lists.


for item in SoC_info:
    write_p(document,item)

for item in Pattern_info:
    write_p(document,item)

# %% Save document
document.save('Scan_items.docx')